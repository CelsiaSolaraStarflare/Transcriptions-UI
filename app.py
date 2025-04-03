import streamlit as st
import dashscope
from dashscope.audio.asr import Recognition
from http import HTTPStatus
from docx import Document
import os

# Set up API Key
DASHSCOPE_API_KEY = "sk-e8dfa404853d43e9870570c6c98c9516"  # Replace with your actual API key
dashscope.api_key = DASHSCOPE_API_KEY

CACHE_DIR = "transcription_cache"
os.makedirs(CACHE_DIR, exist_ok=True)

def transcribe_audio(audio_path):
    recognition = Recognition(
        model='paraformer-realtime-v2',
        format='wav',  # Ensure uploaded files are in WAV format
        sample_rate=16000,
        language_hints=['zh', 'en'],
        callback=None
    )
    result = recognition.call(audio_path)
    
    if result.status_code == HTTPStatus.OK:
        return result.get_sentence()
    else:
        return None

def save_transcription(transcript, cache_subdir):
    output_path = os.path.join(cache_subdir, "transcript.docx")
    doc = Document()
    doc.add_heading("语音转录结果", level=1)
    doc.add_paragraph(transcript)
    doc.save(output_path)
    return output_path

st.title("语音转录系统 (DashScope ASR API)")
st.write("上传WAV音频文件，生成带时间戳的Word文档。")

uploaded_file = st.file_uploader("选择音频文件 (仅支持 WAV 格式)", type=["wav"])

if uploaded_file:
    task_id = str(len(os.listdir(CACHE_DIR)) + 1)
    cache_subdir = os.path.join(CACHE_DIR, task_id)
    os.makedirs(cache_subdir, exist_ok=True)
    
    audio_path = os.path.join(cache_subdir, uploaded_file.name)
    with open(audio_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    
    transcript = transcribe_audio(audio_path)
    
    if transcript:
        doc_path = save_transcription(transcript, cache_subdir)
        
        st.download_button(
            "下载完整DOCX",
            open(doc_path, "rb").read(),
            file_name="full_transcript.docx"
        )
        st.text_area("转录预览", value=transcript, height=300)
    else:
        st.error("转录失败，请检查音频文件或API设置。")
